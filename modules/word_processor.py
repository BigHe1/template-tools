"""
合同模板 — Word 文档处理引擎
查找替换占位符 ▦▦ ?? → ▦数字▦，保留图片
"""

import os
import re
import copy
import io
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class WordProcessor:
    """Word 文档占位符替换处理器"""

    def __init__(self):
        self.ph1 = "▦▦"
        self.ph2 = "??"
        self.counter = 0
        self.checkbox_numbers = []
        self.all_numbers = []

    def set_placeholders(self, ph1: str, ph2: str):
        """设置占位符"""
        if ph1:
            self.ph1 = ph1
        if ph2:
            self.ph2 = ph2

    def process_from_bytes(self, file_bytes: bytes, ph1: str = "", ph2: str = "") -> dict:
        """
        从文件字节流处理 Word 文档

        返回: { "status": "ok"|"error", "result_bytes": bytes, "message": str,
                 "checkbox_numbers": [...], "total": int, "filename": str }
        """
        if not HAS_DOCX:
            return {"status": "error", "message": "缺少 python-docx 依赖，请运行: pip install python-docx"}

        self.set_placeholders(ph1, ph2)
        self._reset()

        try:
            doc = Document(io.BytesIO(file_bytes))

            total = self._replace_in_doc(doc)
            if total == 0:
                return {"status": "error", "message": f"未找到任何占位符（{self.ph1} 或 {self.ph2}）"}

            # 保存到 BytesIO
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            return {
                "status": "ok",
                "result_bytes": output.getvalue(),
                "total": total,
                "checkbox_numbers": self.checkbox_numbers,
                "all_numbers": self.all_numbers,
                "message": f"替换完成，共 {total} 处；复选框位置：{','.join(map(str, self.checkbox_numbers)) if self.checkbox_numbers else '无'}",
            }

        except Exception as e:
            return {"status": "error", "message": f"处理失败：{str(e)}"}

    def _reset(self):
        self.counter = 0
        self.checkbox_numbers.clear()
        self.all_numbers.clear()

    def _replace_in_doc(self, doc) -> int:
        count = 0
        for para in doc.paragraphs:
            count += self._safe_replace_paragraph(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        count += self._safe_replace_paragraph(para)
        return count

    def _safe_replace_paragraph(self, paragraph) -> int:
        total = 0
        runs_to_replace = list(paragraph.runs)
        for run in runs_to_replace:
            total += self._safe_replace_run(run)
        return total

    def _run_has_picture(self, run) -> bool:
        if run._element is None:
            return False
        for child in run._element.iterchildren():
            if child.tag == qn('w:drawing'):
                return True
        return False

    def _safe_replace_run(self, run) -> int:
        text = run.text
        if self.ph1 not in text and self.ph2 not in text:
            return 0

        has_picture = self._run_has_picture(run)

        if not has_picture:
            return self._replace_text_in_run(run, text)

        return self._split_run_with_pictures(run, text)

    def _replace_text_in_run(self, run, original_text) -> int:
        cnt = 0
        new_text = original_text

        while self.ph1 in new_text:
            self.counter += 1
            new_text = new_text.replace(self.ph1, f"▦{self.counter}▦", 1)
            self.all_numbers.append(self.counter)
            cnt += 1

        while self.ph2 in new_text:
            self.counter += 1
            new_text = new_text.replace(self.ph2, f"▦{self.counter}▦", 1)
            self.all_numbers.append(self.counter)
            self.checkbox_numbers.append(self.counter)
            cnt += 1

        if new_text != original_text:
            run.text = new_text
        return cnt

    def _split_run_with_pictures(self, run, original_text) -> int:
        r_element = run._element
        parent_element = run._parent._element

        rPr = None
        rPr_tag = qn('w:rPr')
        for child in r_element.iterchildren():
            if child.tag == rPr_tag:
                rPr = copy.deepcopy(child)
                break

        new_runs_elements = []
        text_buffer = []

        for child in r_element.iterchildren():
            if child.tag == qn('w:t'):
                if child.text:
                    text_buffer.append(child.text)
            elif child.tag == qn('w:drawing'):
                if text_buffer:
                    combined_text = ''.join(text_buffer)
                    new_text = self._process_text_with_placeholders(combined_text)
                    if new_text:
                        text_run = self._create_new_run_element(rPr, new_text)
                        new_runs_elements.append(text_run)
                    text_buffer.clear()
                picture_run = self._create_picture_run_element(rPr, child)
                new_runs_elements.append(picture_run)

        if text_buffer:
            combined_text = ''.join(text_buffer)
            new_text = self._process_text_with_placeholders(combined_text)
            if new_text:
                text_run = self._create_new_run_element(rPr, new_text)
                new_runs_elements.append(text_run)

        if not new_runs_elements:
            return 0

        idx = parent_element.index(r_element)
        for i, new_r in enumerate(new_runs_elements):
            parent_element.insert(idx + i, new_r)
        parent_element.remove(r_element)

        return original_text.count(self.ph1) + original_text.count(self.ph2)

    def _process_text_with_placeholders(self, text: str) -> str:
        new_text = text
        while self.ph1 in new_text:
            self.counter += 1
            new_text = new_text.replace(self.ph1, f"▦{self.counter}▦", 1)
            self.all_numbers.append(self.counter)
        while self.ph2 in new_text:
            self.counter += 1
            new_text = new_text.replace(self.ph2, f"▦{self.counter}▦", 1)
            self.all_numbers.append(self.counter)
            self.checkbox_numbers.append(self.counter)
        return new_text

    def _create_new_run_element(self, rPr, text: str):
        new_run = OxmlElement('w:r')
        if rPr is not None:
            new_run.append(copy.deepcopy(rPr))
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        return new_run

    def _create_picture_run_element(self, rPr, drawing_element):
        new_run = OxmlElement('w:r')
        if rPr is not None:
            new_run.append(copy.deepcopy(rPr))
        new_run.append(copy.deepcopy(drawing_element))
        return new_run
